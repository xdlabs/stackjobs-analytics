import re 
import docx

#the required pattern is stored in p

p=re.compile(r'(.*?)([a-z]+@[a-z]+\.[a-z]{2,4})')

#opening the document

doc=docx.Document('resume.docx')

#assuming the first paragraph always contains the name

name=doc.paragraphs[0].text

#empty list for address

l=[]

#extracting email and address from remaining resume

for i in range(1,len(doc.paragraphs)):
    m=p.search(doc.paragraphs[i].text)
    if m!=None:
        email=m.group(2)
        l.append(m.group(1))
	address=' '.join(l)
        break
    else:
        l.append(doc.paragraphs[i].text)
	address=' '.join(l)

#email contains the email id
#address contains the whole address